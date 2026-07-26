import hashlib
import io
import mimetypes
import uuid

from bs4 import BeautifulSoup
from docx import Document as WordDocument
from pypdf import PdfReader
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from agent_yhzh.config import settings
from agent_yhzh.models import (
    Document,
    DocumentChunk,
    ImportJob,
    KnowledgeCandidate,
    KnowledgeEmbedding,
    KnowledgeEvidence,
)
from agent_yhzh.observability import IMPORT_JOBS
from agent_yhzh.security import CallerContext
from agent_yhzh.services.embeddings import content_hash, embed_text
from agent_yhzh.services.model_config import get_runtime_model_config
from agent_yhzh.services.object_store import get_object, put_object
from agent_yhzh.services.taxonomy import classify_text_llm


ALLOWED_SUFFIXES = {".txt", ".md", ".html", ".htm", ".pdf", ".docx"}


def _suffix(filename: str) -> str:
    return "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def parse_document(data: bytes, filename: str, mime_type: str) -> str:
    suffix = _suffix(filename)
    if suffix not in ALLOWED_SUFFIXES:
        raise ValueError("unsupported_document_type")
    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        document = WordDocument(io.BytesIO(data))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
    text = data.decode("utf-8", errors="replace")
    if suffix in {".html", ".htm"} or "html" in mime_type:
        return BeautifulSoup(text, "html.parser").get_text("\n")
    return text


def chunk_text(content: str) -> list[str]:
    normalized = "\n".join(line.strip() for line in content.splitlines())
    normalized = "\n".join(line for line in normalized.splitlines() if line)
    size = settings.document_chunk_chars
    overlap = min(settings.document_chunk_overlap, size // 2)
    chunks: list[str] = []
    start = 0
    while start < len(normalized):
        end = min(len(normalized), start + size)
        chunk = normalized[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(normalized):
            break
        start = end - overlap
    return chunks


async def create_document_upload(
    session: AsyncSession,
    context: CallerContext,
    *,
    filename: str,
    mime_type: str | None,
    data: bytes,
) -> tuple[Document, ImportJob]:
    if not data or len(data) > settings.document_max_bytes:
        raise ValueError("invalid_document_size")
    if _suffix(filename) not in ALLOWED_SUFFIXES:
        raise ValueError("unsupported_document_type")
    checksum = hashlib.sha256(data).hexdigest()
    existing = await session.scalar(
        select(Document).where(
            Document.tenant_id == context.tenant_id,
            Document.space_id == context.space_id,
            Document.checksum == checksum,
        )
    )
    if existing:
        job = await session.scalar(
            select(ImportJob)
            .where(ImportJob.document_id == existing.id)
            .order_by(ImportJob.created_at.desc())
        )
        if job is None:
            job = ImportJob(
                tenant_id=context.tenant_id,
                space_id=context.space_id,
                document_id=existing.id,
                status="completed",
                progress=100,
            )
            session.add(job)
            await session.commit()
            await session.refresh(job)
        return existing, job

    mime = mime_type or mimetypes.guess_type(filename)[0] or "application/octet-stream"
    document_id = uuid.uuid4()
    object_key = f"{context.tenant_id}/{context.space_id}/{document_id}/{filename}"
    await put_object(object_key, data, mime)
    document = Document(
        id=document_id,
        tenant_id=context.tenant_id,
        space_id=context.space_id,
        filename=filename[:320],
        object_key=object_key,
        checksum=checksum,
        mime_type=mime,
        byte_size=len(data),
    )
    session.add(document)
    await session.flush()
    job = ImportJob(
        tenant_id=context.tenant_id,
        space_id=context.space_id,
        document_id=document.id,
    )
    session.add(job)
    await session.commit()
    await session.refresh(document)
    await session.refresh(job)
    IMPORT_JOBS.labels(status="queued").inc()
    return document, job


async def process_document_import(session: AsyncSession, job_id: uuid.UUID) -> None:
    job = await session.scalar(
        select(ImportJob).where(ImportJob.id == job_id).with_for_update()
    )
    if job is None or job.status == "completed":
        return
    document = await session.get(Document, job.document_id)
    if document is None:
        return
    try:
        job.status = "processing"
        job.progress = 5
        document.parser_status = "processing"
        await session.commit()
        data = await get_object(document.object_key)
        content = parse_document(data, document.filename, document.mime_type)
        chunks = chunk_text(content)
        if not chunks:
            raise ValueError("document_contains_no_text")

        runtime = await get_runtime_model_config(
            session, document.tenant_id, document.space_id
        )
        for index, chunk_content in enumerate(chunks):
            chunk = DocumentChunk(
                tenant_id=document.tenant_id,
                space_id=document.space_id,
                document_id=document.id,
                content=chunk_content,
                location={"chunk": index + 1, "total": len(chunks)},
                token_count=max(1, len(chunk_content) // 3),
                content_hash=content_hash(chunk_content),
            )
            session.add(chunk)
            await session.flush()
            vector = await embed_text(chunk_content, runtime)
            session.add(
                KnowledgeEmbedding(
                    tenant_id=document.tenant_id,
                    space_id=document.space_id,
                    object_type="document_chunk",
                    object_id=chunk.id,
                    model=runtime.embedding_model or settings.embedding_model,
                    vector=vector,
                    content_hash=chunk.content_hash,
                )
            )
            candidate = KnowledgeCandidate(
                tenant_id=document.tenant_id,
                space_id=document.space_id,
                normalized_key=f"document:{document.id}:{index}",
                title=f"{document.filename} · 第 {index + 1} 段",
                content=chunk_content,
                candidate_type="document",
                category=await classify_text_llm(chunk_content, runtime),
                status="pending_review",
                occurrence_count=1,
                distinct_user_count=0,
                score=0.65,
                source_chunk_ids=[str(chunk.id)],
                evidence_summary=f"由文档 {document.filename} 导入，需管理员核验后发布。",
            )
            session.add(candidate)
            await session.flush()
            session.add(
                KnowledgeEvidence(
                    tenant_id=document.tenant_id,
                    space_id=document.space_id,
                    candidate_id=candidate.id,
                    chunk_id=chunk.id,
                    source_kind="document",
                    quote=chunk_content[:5000],
                    confidence=0.8,
                    properties={"filename": document.filename, "chunk": index + 1},
                )
            )
            job.progress = min(95, 10 + int((index + 1) / len(chunks) * 85))
            await session.flush()
        document.parser_status = "completed"
        job.status = "completed"
        job.progress = 100
        await session.commit()
        IMPORT_JOBS.labels(status="completed").inc()
    except Exception as error:
        await session.rollback()
        job = await session.get(ImportJob, job_id)
        document = await session.get(Document, job.document_id) if job else None
        if job:
            job.status = "failed"
            job.error = str(error)[:4000]
        if document:
            document.parser_status = "failed"
        await session.commit()
        IMPORT_JOBS.labels(status="failed").inc()
        raise


async def list_documents(
    session: AsyncSession, tenant_id: str, space_id: str, limit: int = 100
) -> list[Document]:
    return list(
        await session.scalars(
            select(Document)
            .where(Document.tenant_id == tenant_id, Document.space_id == space_id)
            .order_by(Document.created_at.desc())
            .limit(limit)
        )
    )


async def list_import_jobs(
    session: AsyncSession, tenant_id: str, space_id: str, limit: int = 100
) -> list[ImportJob]:
    return list(
        await session.scalars(
            select(ImportJob)
            .where(ImportJob.tenant_id == tenant_id, ImportJob.space_id == space_id)
            .order_by(ImportJob.created_at.desc())
            .limit(limit)
        )
    )
